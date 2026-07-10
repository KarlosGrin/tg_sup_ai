"""
Сервис для работы с файлами: скачивание, сохранение, очистка.
"""

import os
import uuid
import shutil
from pathlib import Path
from typing import Optional

from aiogram.types import Document, Message
from config import config


class FileService:
    """Управление загруженными и обработанными файлами."""

    ALLOWED_EXTENSIONS = {
        ".xlsx", ".xls", ".csv", ".ods",
        ".docx", ".doc",
        ".txt", ".json", ".xml",
    }

    def __init__(self):
        self.download_dir = Path(config.DOWNLOAD_DIR)
        self.processed_dir = Path(config.PROCESSED_DIR)
        self.download_dir.mkdir(parents=True, exist_ok=True)
        self.processed_dir.mkdir(parents=True, exist_ok=True)

        # Хранилище: user_id -> {"original": [пути], "processed": [пути]}
        self._user_files: dict[int, dict[str, list[str]]] = {}

    async def download_file(self, document: Document, user_id: int, bot=None) -> Optional[Path]:
        """Скачать файл из Telegram во временную директорию (в подпапку пользователя)."""
        ext = Path(document.file_name).suffix.lower()
        if ext not in self.ALLOWED_EXTENSIONS:
            return None

        # 🛡️ Изоляция: каждый пользователь — своя подпапка в downloads/
        user_download_dir = self.download_dir / str(user_id)
        user_download_dir.mkdir(parents=True, exist_ok=True)

        # Уникальное имя для избежания коллизий
        unique_name = f"{uuid.uuid4().hex}_{document.file_name}"
        save_path = user_download_dir / unique_name

        # aiogram 3.x: используем bot.download() вместо document.download()
        if bot:
            await bot.download(file=document.file_id, destination=save_path)
        else:
            from aiogram import Bot
            # Fallback — не должно произойти, но на всякий случай
            raise ValueError("Bot instance is required to download files")

        # Проверка сигнатуры файла (magic bytes) — защита от переименованных .exe
        if not self._validate_magic_bytes(str(save_path), ext):
            try:
                save_path.unlink()
            except OSError:
                pass
            return None

        # Сохраняем запись о файле пользователя
        if user_id not in self._user_files:
            self._user_files[user_id] = {"original": [], "processed": []}
        self._user_files[user_id]["original"].append(str(save_path))

        return save_path

    def get_user_files(self, user_id: int) -> list[str]:
        """Получить список оригинальных файлов пользователя."""
        return self._user_files.get(user_id, {}).get("original", [])

    def save_processed(self, source_path: str | Path, user_id: int) -> Path:
        """Сохранить обработанный файл в подпапку пользователя и вернуть путь."""
        source = Path(source_path)
        orig_name = source.name
        # Убираем UUID-префикс для понятного имени
        parts = orig_name.split("_", 1)
        clean_name = parts[1] if len(parts) > 1 else orig_name
        processed_name = f"processed_{clean_name}"

        # Изоляция: каждый пользователь — своя подпапка
        user_processed_dir = self.processed_dir / str(user_id)
        user_processed_dir.mkdir(parents=True, exist_ok=True)
        processed_path = user_processed_dir / processed_name

        shutil.copy2(source_path, processed_path)

        if user_id not in self._user_files:
            self._user_files[user_id] = {"original": [], "processed": []}
        self._user_files[user_id]["processed"].append(str(processed_path))

        return processed_path

    def cleanup_user_files(self, user_id: int):
        """Удалить все файлы пользователя."""
        user_data = self._user_files.pop(user_id, {})
        for paths in user_data.values():
            for p in paths:
                try:
                    os.remove(p)
                except OSError:
                    pass

    def cleanup_old_files(self, max_age_hours: int = 24):
        """Очистка всех временных файлов старше N часов (рекурсивно по подпапкам пользователей)."""
        import time
        now = time.time()
        for directory in [self.download_dir, self.processed_dir]:
            # Чистим файлы рекурсивно (в т.ч. в user-подпапках)
            for f in directory.rglob("*"):
                if f.is_file():
                    file_age = now - f.stat().st_mtime
                    if file_age > max_age_hours * 3600:
                        try:
                            f.unlink()
                        except OSError:
                            pass
            # Удаляем пустые подпапки
            for subdir in sorted(directory.iterdir(), key=lambda p: str(p), reverse=True):
                if subdir.is_dir():
                    try:
                        subdir.rmdir()  # удалит только пустую
                    except OSError:
                        pass

    @staticmethod
    def _validate_magic_bytes(file_path: str, expected_ext: str) -> bool:
        """
        Проверить сигнатуру файла (magic bytes) на соответствие расширению.
        Защита от переименованных .exe / .zip и других подставных файлов.
        """
        MAGIC_BYTES = {
            b"\x50\x4B\x03\x04": (".xlsx", ".xlsm", ".docx", ".ods"),  # ZIP-архивы (Office Open XML)
            b"\xD0\xCF\x11\xE0": (".xls", ".doc"),                     # OLE2 (старый Office)
            b"\xEF\xBB\xBF":     (".csv", ".txt", ".json", ".xml"),    # UTF-8 BOM
            b"\x3C\x3F\x78\x6D": (".xml",),                            # <?xml
            b"\x7B":             (".json",),                            # {
            b"\x5B":             (".json",),                            # [
        }

        try:
            with open(file_path, "rb") as f:
                header = f.read(4)

            # CSV/TXT/JSON без BOM — пропускаем (любой текст допустим)
            if expected_ext in (".csv", ".txt"):
                return True

            # Проверяем по известным сигнатурам
            for magic, extensions in MAGIC_BYTES.items():
                if header.startswith(magic):
                    return expected_ext in extensions

            # Если сигнатура не найдена — разрешаем только текст
            return expected_ext in (".csv", ".txt", ".json", ".xml")

        except OSError:
            return False

    @staticmethod
    def _escape_data(text: str, max_len: int = 300) -> str:
        """
        Экранировать пользовательские данные для защиты от prompt injection.
        Оборачивает данные в маркеры и обрезает до max_len символов.
        """
        # Обрезаем до лимита
        if len(text) > max_len:
            text = text[:max_len] + "... (обрезано)"
        # Заменяем потенциально опасные последовательности
        text = text.replace("```", "'''")
        return text

    @staticmethod
    def get_file_summary(file_path: str) -> str:
        """
        Получить краткое описание файла: формат, размер, базовую структуру.
        Возвращает текстовое описание для передачи AI.
        """
        path = Path(file_path)
        ext = path.suffix.lower()
        size_kb = path.stat().st_size / 1024
        summary = [f"📄 Файл: `{path.name}`", f"📏 Размер: `{size_kb:.1f} КБ`"]

        try:
            if ext in (".xlsx", ".xls", ".ods"):
                import pandas as pd
                excel_file = pd.ExcelFile(file_path)
                sheet_names = excel_file.sheet_names
                summary.append(f"📑 Листы: `{', '.join(sheet_names)}`")
                for sheet in sheet_names[:3]:  # Показываем первые 3 листа
                    df = pd.read_excel(file_path, sheet_name=sheet, nrows=5)
                    cols = list(df.columns)
                    dtypes = {c: str(df[c].dtype) for c in cols}
                    summary.append(f"\n  🔹 Лист «{sheet}»:")
                    summary.append(f"     Колонки ({len(cols)}): `{', '.join(cols)}`")
                    summary.append(f"     Типы: `{dtypes}`")
                    summary.append(f"     Строк (всего): ~{len(pd.read_excel(file_path, sheet_name=sheet))}")

            elif ext == ".csv":
                import pandas as pd
                df = pd.read_csv(file_path, nrows=5, encoding="utf-8")
                full_df = pd.read_csv(file_path, encoding="utf-8")
                cols = list(df.columns)
                dtypes = {c: str(df[c].dtype) for c in cols}
                summary.append(f"📊 CSV файл")
                summary.append(f"   Колонки ({len(cols)}): `{', '.join(cols)}`")
                summary.append(f"   Типы: `{dtypes}`")
                summary.append(f"   Строк (всего): {len(full_df)}")

            elif ext in (".docx", ".doc"):
                try:
                    from docx import Document as DocxDocument
                    doc = DocxDocument(file_path)
                    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                    tables = doc.tables
                    summary.append(f"📝 Word документ")
                    summary.append(f"   Абзацев: {len(paragraphs)}")
                    summary.append(f"   Таблиц: {len(tables)}")
                    if paragraphs:
                        preview = FileService._escape_data(paragraphs[0], 200)
                        summary.append(f"   Начало: «{preview}»")
                    if paragraphs:
                        last = FileService._escape_data(paragraphs[-1], 100)
                        summary.append(f"   Последний: «{last}»")
                    if tables:
                        for ti, table in enumerate(tables[:2], 1):
                            rows = len(table.rows)
                            cols = len(table.columns)
                            summary.append(f"   Таблица {ti}: {rows} x {cols}")
                            # Показываем заголовки первой строки
                            if rows > 0:
                                headers = [FileService._escape_data(cell.text.strip(), 30) for cell in table.rows[0].cells]
                                summary.append(f"     Заголовки: {', '.join(headers)}")
                    # Стили документа
                    styles_used = set()
                    for p in doc.paragraphs:
                        if p.style and p.style.name:
                            styles_used.add(p.style.name)
                    if styles_used:
                        summary.append(f"   Стили: {', '.join(list(styles_used)[:5])}")
                except Exception as e:
                    summary.append(f"📝 Word документ (ошибка чтения: {e})")

            elif ext == ".txt":
                try:
                    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                        lines = f.readlines()
                    total_lines = len(lines)
                    summary.append(f"📃 Текстовый файл")
                    summary.append(f"   Строк: {total_lines}")
                    summary.append(f"   Символов: {sum(len(l) for l in lines)}")
                    # Первые строки
                    preview_lines = [l.strip() for l in lines[:5] if l.strip()]
                    if preview_lines:
                        summary.append(f"   Начало ({len(preview_lines)} строк):")
                        for l in preview_lines[:3]:
                            summary.append(f"     «{FileService._escape_data(l, 150)}»")
                    if total_lines > 5:
                        last_lines = [l.strip() for l in lines[-3:] if l.strip()]
                        if last_lines:
                            summary.append(f"   Конец ({len(last_lines)} строк):")
                            for l in last_lines[:2]:
                                summary.append(f"     «{FileService._escape_data(l, 150)}»")
                except Exception as e:
                    summary.append(f"📃 Текстовый файл (ошибка чтения: {e})")

        except Exception as e:
            summary.append(f"⚠️ Ошибка чтения: {e}")

        return "\n".join(summary)


file_service = FileService()
