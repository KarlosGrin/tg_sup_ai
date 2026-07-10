"""
Unit-тесты для CodeExecutor (песочница исполнения кода).
Критически важны для безопасности — проверяют, что sandbox блокирует опасные операции.

Запуск:
    python -m pytest tests/test_code_executor.py -v --tb=short
"""

import pytest
import json
import sys
from pathlib import Path

# Добавляем корень проекта в sys.path для импорта сервисов
sys.path.insert(0, str(Path(__file__).parent.parent))

from services.code_executor import CodeExecutor, _BLOCKED_PATTERNS


@pytest.fixture(autouse=True)
def _patch_docker_disabled(monkeypatch):
    """Принудительно отключаем Docker для всех тестов (используем прямой subprocess).
    В .env может быть DOCKER_ENABLED=true, это ломает тесты.
    """
    monkeypatch.setattr("config.config.DOCKER_ENABLED", False)
    monkeypatch.setattr("config.config.EXECUTION_TIMEOUT", 30)  # быстрый таймаут для тестов


@pytest.fixture
def executor(tmp_path, monkeypatch):
    """CodeExecutor с временными директориями для каждого теста."""
    download_dir = tmp_path / "downloads"
    processed_dir = tmp_path / "processed"
    download_dir.mkdir()
    processed_dir.mkdir()

    ex = CodeExecutor()
    monkeypatch.setattr(ex, "_download_dir", str(download_dir.resolve()))
    monkeypatch.setattr(ex, "_processed_dir", str(processed_dir.resolve()))
    return ex


@pytest.fixture
def sample_excel(tmp_path):
    """Создать тестовый Excel файл."""
    download_dir = tmp_path / "downloads"
    download_dir.mkdir(exist_ok=True)
    import pandas as pd
    df = pd.DataFrame({
        "name": ["Alice", "Bob", "Charlie"],
        "age": [25, 30, 35],
        "salary": [50000, 60000, 70000],
    })
    file_path = download_dir / "test_data.xlsx"
    df.to_excel(file_path, index=False)
    return file_path


@pytest.fixture
def output_path(tmp_path):
    """Путь для сохранения результата."""
    processed_dir = tmp_path / "processed"
    processed_dir.mkdir(exist_ok=True)
    return processed_dir / "result.xlsx"


# ═══════════════════════════════════════════════════════════════
# Блокировка опасных импортов
# ═══════════════════════════════════════════════════════════════

class TestBlockDangerousImports:
    """Тесты: опасные модули должны быть заблокированы."""

    @pytest.mark.parametrize("module", [
        "os", "subprocess", "socket", "ctypes",
        "shutil", "multiprocessing", "threading",
        "requests", "urllib", "httpx", "aiohttp",
        "pickle", "marshal", "base64",
        "cryptography", "signal",
    ])
    def test_block_module_import(self, module, executor, sample_excel, output_path):
        """Модуль {module} должен быть заблокирован sandbox-ом."""
        code = f"import {module}"
        result = executor.execute(code, str(sample_excel), str(output_path))
        assert not result["success"], f"Module '{module}' should be blocked"
        assert module in result["stderr"] or "запрещён" in result["stderr"]

    @pytest.mark.parametrize("pattern", [
        "eval(", "exec(", "compile(",
    ])
    def test_block_dangerous_patterns_before_sandbox(self, pattern, executor, sample_excel, output_path):
        """Паттерн {pattern} должен быть заблокирован ДО sandbox (_check_blocked_patterns)."""
        code = f"x = {pattern}1+1')"
        result = executor.execute(code, str(sample_excel), str(output_path))
        assert not result["success"]
        assert pattern.replace("(", "") in result["stderr"] or "Блокирован" in result["stderr"]

    @pytest.mark.parametrize("pattern", [
        "__reduce__", "__reduce_ex__", "__mro__", "__subclasses__",
        "__globals__", "__code__", "__builtins__",
    ])
    def test_block_dunder_patterns(self, pattern, executor, sample_excel, output_path):
        """Дандер-паттерн {pattern} должен быть заблокирован."""
        code = f"x = [].{pattern}"
        result = executor.execute(code, str(sample_excel), str(output_path))
        assert not result["success"]
        assert pattern in result["stderr"]

    def test_block_breakpoint(self, executor, sample_excel, output_path):
        """breakpoint() должен быть заблокирован."""
        code = "breakpoint()"
        result = executor.execute(code, str(sample_excel), str(output_path))
        assert not result["success"]
        assert "breakpoint" in result["stderr"]


# ═══════════════════════════════════════════════════════════════
# AST-анализ: обход блок-листа через конкатенацию строк
# ═══════════════════════════════════════════════════════════════

class TestASTAnalysis:
    """Тесты: AST-анализ должен ловить обфусцированный код (конкатенация строк)."""

    # Прямое обращение через .__subclasses__() должно блокироваться AST
    @pytest.mark.parametrize("attr", [
        "__subclasses__", "__bases__", "__mro__",
        "__globals__", "__code__", "__reduce__", "__reduce_ex__",
    ])
    def test_block_dunder_attribute_direct(self, attr, executor, sample_excel, output_path):
        """Прямой вызов .{attr}() должен блокироваться AST."""
        code = f"x = ().__class__.__bases__[0].{attr}()"
        result = executor.execute(code, str(sample_excel), str(output_path))
        assert not result["success"]
        assert "AST" in result["stderr"] or attr in result["stderr"]

    # getattr(base, 'строка') — обход через разбивку
    @pytest.mark.parametrize("attr", [
        "__subclasses__", "__bases__", "__mro__",
        "__globals__", "__code__",
    ])
    def test_block_getattr_string(self, attr, executor, sample_excel, output_path):
        """getattr(obj, '{attr}') должен блокироваться AST."""
        code = f"x = getattr(().__class__.__bases__[0], '{attr}')"
        result = executor.execute(code, str(sample_excel), str(output_path))
        assert not result["success"]
        assert "AST" in result["stderr"] or attr in result["stderr"]

    # Конкатенация строковых литералов: '__subcla' + 'sses__'
    @pytest.mark.parametrize("parts,dunder", [
        (["__subcla", "sses__"], "__subclasses__"),
        (["__ba", "ses__"], "__bases__"),
        (["__m", "ro__"], "__mro__"),
        (["__glo", "bals__"], "__globals__"),
        (["__cod", "e__"], "__code__"),
        (["__reduc", "e__"], "__reduce__"),
        (["__reduce_", "ex__"], "__reduce_ex__"),
    ])
    def test_block_string_concatenation(self, parts, dunder, executor, sample_excel, output_path):
        """Конкатенация строк '{parts[0]}' + '{parts[1]}' должна блокироваться AST."""
        code = f"""
base = ().__class__.__bases__[0]
name = '{parts[0]}' '{parts[1]}'
for cls in getattr(base, name)():
    if cls.__name__ == 'Popen':
        pass
"""
        result = executor.execute(code, str(sample_excel), str(output_path))
        assert not result["success"]
        assert "AST" in result["stderr"] or "запрещён" in result["stderr"]

    # __getattribute__('__subclasses__')
    @pytest.mark.parametrize("attr", [
        "__subclasses__", "__bases__", "__mro__",
        "__globals__", "__code__",
    ])
    def test_block_getattribute(self, attr, executor, sample_excel, output_path):
        """obj.__getattribute__('{attr}') должен блокироваться AST."""
        code = f"x = object.__getattribute__(object(), '{attr}')"
        result = executor.execute(code, str(sample_excel), str(output_path))
        assert not result["success"]
        assert "AST" in result["stderr"] or attr in result["stderr"]

    # Легитимный код не должен ложно срабатывать
    def test_ast_allows_legitimate_getattr(self, executor, sample_excel, output_path):
        """getattr с обычными именами атрибутов должен работать."""
        code = """
import pandas as pd
df = pd.read_excel(input_path)
col_name = getattr(df, 'name')
print(f"Column type: {type(col_name)}")
df.to_excel(output_path, index=False)
"""
        result = executor.execute(code, str(sample_excel), str(output_path))
        assert result["success"], f"Legitimate getattr blocked: {result['stderr']}"
        assert output_path.exists()

    def test_ast_allows_class_name_access(self, executor, sample_excel, output_path):
        """obj.__class__.__name__ — легитимный паттерн, не должен блокироваться."""
        code = """
import pandas as pd
df = pd.read_excel(input_path)
print(f"Type: {df['name'].__class__.__name__}")
df.to_excel(output_path, index=False)
"""
        result = executor.execute(code, str(sample_excel), str(output_path))
        assert result["success"], f"__class__.__name__ blocked: {result['stderr']}"


# ═══════════════════════════════════════════════════════════════
# Изоляция файловой системы
# ═══════════════════════════════════════════════════════════════

class TestFileSystemIsolation:
    """Тесты: файловая система должна быть изолирована."""

    def test_block_write_outside(self, executor, sample_excel, output_path):
        """Запись вне разрешённых директорий должна быть заблокирована."""
        code = """
with open('C:\\\\Windows\\\\win.ini', 'w') as f:
    f.write('hacked')
"""
        result = executor.execute(code, str(sample_excel), str(output_path))
        assert not result["success"]
        assert "Access denied" in result["stderr"] or "PermissionError" in result["stderr"]

    def test_block_read_outside(self, executor, sample_excel, output_path):
        """Чтение вне разрешённых директорий должно быть заблокировано."""
        code = """
with open('C:\\\\Windows\\\\win.ini', 'r') as f:
    content = f.read()
"""
        result = executor.execute(code, str(sample_excel), str(output_path))
        assert not result["success"]
        assert "Access denied" in result["stderr"] or "PermissionError" in result["stderr"]

    def test_allow_write_to_allowed_dir(self, executor, sample_excel, tmp_path):
        """Запись внутри allowed_dir должна работать."""
        processed_dir = tmp_path / "processed"
        out_path = processed_dir / "result.txt"
        code = "with open(output_path, 'w') as f: f.write('ok')"
        result = executor.execute(code, str(sample_excel), str(out_path))
        assert result["success"]
        assert out_path.exists()
        assert out_path.read_text() == "ok"


# ═══════════════════════════════════════════════════════════════
# Разрешённые операции (позитивные тесты)
# ═══════════════════════════════════════════════════════════════

class TestAllowedOperations:
    """Тесты: легитимные операции должны работать."""

    def test_pandas_read_write_excel(self, executor, sample_excel, output_path):
        """Pandas: чтение Excel -> трансформация -> запись."""
        code = """
import pandas as pd
df = pd.read_excel(input_path)
df['age_doubled'] = df['age'] * 2
df.to_excel(output_path, index=False)
print(f"Processed {len(df)} rows")
"""
        result = executor.execute(code, str(sample_excel), str(output_path))
        assert result["success"], f"Expected success, got: {result['stderr']}"
        assert output_path.exists()
        import pandas as pd
        df = pd.read_excel(output_path)
        assert "age_doubled" in df.columns
        assert len(df) == 3
        assert "Processed 3 rows" in result["stdout"]

    def test_numpy_operations(self, executor, sample_excel, output_path):
        """NumPy: математические операции."""
        code = """
import numpy as np
import pandas as pd
df = pd.read_excel(input_path)
df['salary_sqrt'] = np.sqrt(df['salary'])
df.to_excel(output_path, index=False)
"""
        result = executor.execute(code, str(sample_excel), str(output_path))
        assert result["success"], f"Expected success, got: {result['stderr']}"
        assert output_path.exists()

    def test_openpyxl_styling(self, executor, sample_excel, output_path):
        """openpyxl: работа со стилями."""
        code = """
import openpyxl
from openpyxl.styles import Font
wb = openpyxl.load_workbook(input_path)
ws = wb.active
ws['A1'].font = Font(bold=True)
wb.save(output_path)
"""
        result = executor.execute(code, str(sample_excel), str(output_path))
        assert result["success"], f"Expected success, got: {result['stderr']}"
        assert output_path.exists()

    def test_docx_operations(self, executor, tmp_path):
        """python-docx: создание/чтение Word."""
        download_dir = tmp_path / "downloads"
        download_dir.mkdir(exist_ok=True)
        from docx import Document
        doc = Document()
        doc.add_paragraph("Test content")
        input_path = download_dir / "test.docx"
        doc.save(input_path)

        processed_dir = tmp_path / "processed"
        processed_dir.mkdir(exist_ok=True)
        output_path = processed_dir / "result.docx"

        code = """
from docx import Document
doc = Document(input_path)
doc.add_paragraph("Added by sandbox")
doc.save(output_path)
"""
        result = executor.execute(code, str(input_path), str(output_path))
        assert result["success"], f"Expected success, got: {result['stderr']}"
        assert output_path.exists()

    def test_csv_operations(self, executor, tmp_path):
        """CSV: чтение/запись."""
        download_dir = tmp_path / "downloads"
        download_dir.mkdir(exist_ok=True)
        import pandas as pd
        df = pd.DataFrame({"product": ["A", "B"], "price": [10.5, 20.0]})
        input_path = download_dir / "test.csv"
        df.to_csv(input_path, index=False)

        processed_dir = tmp_path / "processed"
        processed_dir.mkdir(exist_ok=True)
        output_path = processed_dir / "result.csv"

        code = """
import pandas as pd
df = pd.read_csv(input_path)
df['discounted'] = df['price'] * 0.9
df.to_csv(output_path, index=False)
"""
        result = executor.execute(code, str(input_path), str(output_path))
        assert result["success"], f"Expected success, got: {result['stderr']}"
        assert output_path.exists()

    def test_json_operations(self, executor, tmp_path):
        """JSON: чтение/запись."""
        download_dir = tmp_path / "downloads"
        download_dir.mkdir(exist_ok=True)
        import json
        input_path = download_dir / "test.json"
        input_path.write_text(json.dumps({"items": [1, 2, 3]}))

        processed_dir = tmp_path / "processed"
        processed_dir.mkdir(exist_ok=True)
        output_path = processed_dir / "result.json"

        code = """
import json
with open(input_path, 'r') as f:
    data = json.load(f)
data['processed'] = True
with open(output_path, 'w') as f:
    json.dump(data, f)
"""
        result = executor.execute(code, str(input_path), str(output_path))
        assert result["success"], f"Expected success, got: {result['stderr']}"
        assert output_path.exists()
        data = json.loads(output_path.read_text())
        assert data["processed"] is True

    def test_stdlib_modules(self, executor, sample_excel, output_path):
        """Стандартные библиотеки (re, datetime, math, collections, itertools, statistics, copy)."""
        code = """
import re, datetime, math, collections, itertools, statistics, copy
import pandas as pd
df = pd.read_excel(input_path)
df['name_clean'] = df['name'].apply(lambda x: re.sub(r'[aeiou]', '', x, flags=re.I))
df.to_excel(output_path, index=False)
"""
        result = executor.execute(code, str(sample_excel), str(output_path))
        assert result["success"], f"Expected success, got: {result['stderr']}"
        assert output_path.exists()

    def test_pathlib_operations(self, executor, sample_excel, output_path):
        """pathlib: работа с путями."""
        code = """
from pathlib import Path
import pandas as pd
df = pd.read_excel(input_path)
p = Path(input_path)
print(f"Input file: {p.name}")
df.to_excel(output_path, index=False)
"""
        result = executor.execute(code, str(sample_excel), str(output_path))
        assert result["success"], f"Expected success, got: {result['stderr']}"
        assert output_path.exists()
        assert "Input file:" in result["stdout"]

    def test_xml_operations(self, executor, tmp_path):
        """xml.etree.ElementTree."""
        download_dir = tmp_path / "downloads"
        download_dir.mkdir(exist_ok=True)
        import xml.etree.ElementTree as ET
        root = ET.Element("root")
        ET.SubElement(root, "item").text = "test"
        input_path = download_dir / "test.xml"
        ET.ElementTree(root).write(input_path)

        processed_dir = tmp_path / "processed"
        processed_dir.mkdir(exist_ok=True)
        output_path = processed_dir / "result.xml"

        code = """
import xml.etree.ElementTree as ET
tree = ET.parse(input_path)
root = tree.getroot()
new_item = ET.SubElement(root, "new_item")
new_item.text = "added"
tree.write(output_path)
"""
        result = executor.execute(code, str(input_path), str(output_path))
        assert result["success"], f"Expected success, got: {result['stderr']}"
        assert output_path.exists()


# ═══════════════════════════════════════════════════════════════
# Санитизация кода
# ═══════════════════════════════════════════════════════════════

class TestCodeSanitization:
    """Тесты: _sanitize_code должен очищать AI-сгенерированный код."""

    def test_removes_hardcoded_paths(self, executor, sample_excel, output_path):
        """Хардкоженные input_path/output_path должны комментироваться."""
        code = """
input_path = '/evil/path.xlsx'
output_path = '/evil/output.xlsx'
import pandas as pd
df = pd.read_excel(input_path)
df.to_excel(output_path, index=False)
"""
        result = executor.execute(code, str(sample_excel), str(output_path))
        assert result["success"], f"Expected success, got: {result['stderr']}"
        assert output_path.exists()

    def test_removes_duplicate_imports(self, executor, sample_excel, output_path):
        """Дублирующие import pandas/numpy должны удаляться."""
        code = """
import pandas as pd
from pandas import DataFrame
import numpy as np
import numpy as np
df = pd.read_excel(input_path)
df.to_excel(output_path, index=False)
"""
        result = executor.execute(code, str(sample_excel), str(output_path))
        assert result["success"], f"Expected success, got: {result['stderr']}"
        assert output_path.exists()


# ═══════════════════════════════════════════════════════════════
# Таймаут
# ═══════════════════════════════════════════════════════════════

class TestTimeout:
    """Тесты: код, превышающий таймаут, должен быть принудительно завершён."""

    def test_infinite_loop_is_killed(self, executor, sample_excel, output_path):
        """Бесконечный цикл должен быть остановлен по таймауту (120 сек)."""
        code = """
while True:
    pass
"""
        result = executor.execute(code, str(sample_excel), str(output_path))
        assert not result["success"]
        assert "Превышено время" in result["stderr"] or "timeout" in result["stderr"].lower()


# ═══════════════════════════════════════════════════════════════
# Обработка ошибок в пользовательском коде
# ═══════════════════════════════════════════════════════════════

class TestUserCodeErrors:
    """Тесты: исключения в пользовательском коде не должны крашить sandbox."""

    def test_exception_captured(self, executor, sample_excel, output_path):
        """Исключения должны попадать в stderr."""
        code = """
import pandas as pd
df = pd.read_excel(input_path)
raise ValueError("Test error from user code")
"""
        result = executor.execute(code, str(sample_excel), str(output_path))
        assert not result["success"]
        assert "ValueError" in result["stderr"]
        assert "Test error" in result["stderr"]

    def test_syntax_error_captured(self, executor, sample_excel, output_path):
        """SyntaxError не должен крашить sandbox."""
        code = """import pandas as pd\ndf = pd.read_excel(input_path\n"""
        result = executor.execute(code, str(sample_excel), str(output_path))
        assert not result["success"]
        assert "SyntaxError" in result["stderr"] or "Error" in result["stderr"]

    def test_keyerror_captured(self, executor, sample_excel, output_path):
        """KeyError (колонка не найдена) должен перехватываться."""
        code = """
import pandas as pd
df = pd.read_excel(input_path)
df['new'] = df['nonexistent_column']
df.to_excel(output_path, index=False)
"""
        result = executor.execute(code, str(sample_excel), str(output_path))
        assert not result["success"]
        assert "KeyError" in result["stderr"] or "nonexistent" in result["stderr"]


# ═══════════════════════════════════════════════════════════════
# Вывод (stdout / stderr)
# ═══════════════════════════════════════════════════════════════

class TestOutput:
    """Тесты: захват stdout/stderr."""

    def test_stdout_captured(self, executor, sample_excel, output_path):
        """print() должен захватываться в result['stdout']."""
        code = """
import pandas as pd
df = pd.read_excel(input_path)
print(f"Rows: {len(df)}")
print(f"Columns: {list(df.columns)}")
df.to_excel(output_path, index=False)
"""
        result = executor.execute(code, str(sample_excel), str(output_path))
        assert result["success"]
        assert "Rows: 3" in result["stdout"]
        assert "Columns:" in result["stdout"]

    def test_stderr_captured_on_error(self, executor, sample_excel, output_path):
        """stderr должен захватывать исключения."""
        code = """
import pandas as pd
df = pd.read_excel(input_path)
raise RuntimeError("Test stderr capture")
"""
        result = executor.execute(code, str(sample_excel), str(output_path))
        assert not result["success"]
        assert "RuntimeError" in result["stderr"]
        assert "Test stderr capture" in result["stderr"]


# ═══════════════════════════════════════════════════════════════
# Полнота _BLOCKED_PATTERNS
# ═══════════════════════════════════════════════════════════════

class TestBlockedPatternsCompleteness:
    """Тесты: _BLOCKED_PATTERNS покрывает известные опасные векторы."""

    def test_not_empty(self):
        assert len(_BLOCKED_PATTERNS) > 0

    def test_critical_patterns_present(self):
        """Критические паттерны (CWE-78, CWE-94, CWE-502, CWE-915)."""
        critical = {
            # RCE / OS command injection
            "subprocess": "CWE-78",
            "shutil": "CWE-78",
            "socket": "CWE-78",
            "ctypes": "CWE-78",
            # Code injection
            "eval(": "CWE-94",
            "exec(": "CWE-94",
            "compile(": "CWE-94",
            # Deserialization RCE
            "pickle": "CWE-502",
            "marshal": "CWE-502",
            # MRO / prototype pollution
            "__reduce__": "CWE-915",
            "__reduce_ex__": "CWE-915",
            "__mro__": "CWE-915",
            "__subclasses__": "CWE-915",
            "__globals__": "CWE-915",
            "__code__": "CWE-915",
            "__builtins__": "CWE-915",
            # Network
            "requests": "Network",
            "urllib": "Network",
            "httpx": "Network",
            "aiohttp": "Network",
            # Misc
            "base64": "Obfuscation",
            "breakpoint": "Debug",
            "cryptography": "Crypto",
            "signal": "Signal",
        }
        for pattern, reason in critical.items():
            assert pattern in _BLOCKED_PATTERNS, \
                f"Missing critical pattern '{pattern}' ({reason})"

    def test_no_false_positives_on_legitimate_code(self):
        """Легитимный код не должен содержать блок-паттернов."""
        code = """
import pandas as pd
import numpy as np
df = pd.DataFrame({'a': [1, 2, 3]})
df['b'] = df['a'] * 2
df.to_excel(output_path, index=False)
print(f"Done: {len(df)} rows")
"""
        for pattern in _BLOCKED_PATTERNS:
            assert pattern not in code, \
                f"False positive: pattern '{pattern}' matches legitimate code"


# ═══════════════════════════════════════════════════════════════
# Структура результата
# ═══════════════════════════════════════════════════════════════

class TestResultStructure:
    """Тесты: формат результата execute()."""

    def test_returns_dict_with_expected_keys(self, executor, sample_excel, output_path):
        """execute() должен возвращать dict с success/stdout/stderr/output_path."""
        code = "import pandas as pd\ndf = pd.read_excel(input_path)\ndf.to_excel(output_path, index=False)"
        result = executor.execute(code, str(sample_excel), str(output_path))
        assert isinstance(result, dict)
        assert "success" in result
        assert "stdout" in result
        assert "stderr" in result
        assert "output_path" in result
        assert result["output_path"] == str(output_path)

    def test_success_creates_output_file(self, executor, sample_excel, output_path):
        """При успехе файл result должен существовать."""
        code = "import pandas as pd\ndf = pd.read_excel(input_path)\ndf.to_excel(output_path, index=False)"
        result = executor.execute(code, str(sample_excel), str(output_path))
        assert result["success"]
        assert output_path.exists()


# ═══════════════════════════════════════════════════════════════
# Интеграционные тесты
# ═══════════════════════════════════════════════════════════════

class TestIntegration:
    """Интеграционные тесты полного цикла обработки."""

    def test_excel_to_csv_conversion(self, executor, tmp_path):
        """Конвертация Excel -> CSV."""
        download_dir = tmp_path / "downloads"
        download_dir.mkdir(exist_ok=True)
        import pandas as pd
        df = pd.DataFrame({"a": [1, 2], "b": [3, 4]})
        input_path = download_dir / "input.xlsx"
        df.to_excel(input_path, index=False)

        processed_dir = tmp_path / "processed"
        processed_dir.mkdir(exist_ok=True)
        output_path = processed_dir / "output.csv"

        code = """
import pandas as pd
df = pd.read_excel(input_path)
df.to_csv(output_path, index=False)
"""
        result = executor.execute(code, str(input_path), str(output_path))
        assert result["success"]
        assert output_path.exists()
        df_out = pd.read_csv(output_path)
        assert len(df_out) == 2

    def test_multi_sheet_excel(self, executor, tmp_path):
        """Обработка многостраничного Excel."""
        download_dir = tmp_path / "downloads"
        download_dir.mkdir(exist_ok=True)
        import pandas as pd
        with pd.ExcelWriter(download_dir / "multi.xlsx") as writer:
            pd.DataFrame({"x": [1, 2]}).to_excel(writer, sheet_name="Sheet1", index=False)
            pd.DataFrame({"y": [3, 4]}).to_excel(writer, sheet_name="Sheet2", index=False)

        processed_dir = tmp_path / "processed"
        processed_dir.mkdir(exist_ok=True)
        output_path = processed_dir / "result.xlsx"

        code = """
import pandas as pd
sheets = pd.read_excel(input_path, sheet_name=None)
combined = pd.concat(sheets.values(), ignore_index=True)
combined.to_excel(output_path, index=False)
"""
        result = executor.execute(code, str(download_dir / "multi.xlsx"), str(output_path))
        assert result["success"]
        df_out = pd.read_excel(output_path)
        assert len(df_out) == 4

    def test_analyze_file_method(self, executor, sample_excel):
        """analyze_file должен возвращать структуру файла."""
        result = executor.analyze_file(str(sample_excel))
        assert result["error"] is None
        assert "sheets" in result["info"]
        assert len(result["info"]["sheets"]) > 0


# ═══════════════════════════════════════════════════════════════
# Запуск
# ═══════════════════════════════════════════════════════════════
if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])